# Design Ref: §2.2 Ingest flow + FR-01..FR-05.
# Plan SC: "Three ingestion paths work end-to-end" — partial: module-2 covers
# the functions; module-3 wires the HTTP layer.

from __future__ import annotations

import asyncio
import unittest
from datetime import datetime, timezone
from io import BytesIO

from app.sermon_review import (
    IngestError,
    build_sermon,
    generate_segment_id,
    ingest_from_docx,
    ingest_from_google_docs,
    ingest_from_paste,
    ingest_from_txt,
    split_korean_text,
)


KOREAN_SAMPLE = (
    "오늘 우리는 하나님의 은혜를 보려고 합니다. "
    "은혜는 단지 좋은 감정이 아닙니다. "
    "그것은 하나님께서 주시는 선물입니다."
)


class SegmentIdGeneratorTests(unittest.TestCase):
    def test_default_width_is_three(self) -> None:
        self.assertEqual(generate_segment_id(1, 5), "S001")
        self.assertEqual(generate_segment_id(10, 5), "S010")
        self.assertEqual(generate_segment_id(123, 999), "S123")

    def test_width_widens_at_thousands(self) -> None:
        self.assertEqual(generate_segment_id(1, 1000), "S0001")
        self.assertEqual(generate_segment_id(1000, 1000), "S1000")

    def test_rejects_zero_or_negative(self) -> None:
        with self.assertRaises(ValueError):
            generate_segment_id(0, 10)
        with self.assertRaises(ValueError):
            generate_segment_id(1, 0)


class SplitKoreanTextTests(unittest.TestCase):
    def test_splits_on_korean_sentence_endings(self) -> None:
        result = split_korean_text(KOREAN_SAMPLE)
        self.assertEqual(len(result), 3)
        self.assertTrue(result[0].endswith("합니다."))

    def test_strips_empty_lines(self) -> None:
        result = split_korean_text("   \n\n\n")
        self.assertEqual(result, [])

    def test_treats_single_newlines_as_soft_wraps(self) -> None:
        text = "오늘 우리는 하나님의\n은혜를 보려고 합니다."
        result = split_korean_text(text)
        self.assertEqual(len(result), 1)
        self.assertIn("은혜를", result[0])

    def test_auto_split_false_uses_lines(self) -> None:
        text = "라인 1\n라인 2\n라인 3"
        result = split_korean_text(text, auto_split=False)
        self.assertEqual(result, ["라인 1", "라인 2", "라인 3"])


class IngestFromPasteTests(unittest.TestCase):
    def test_normalizes_and_returns_text(self) -> None:
        out = ingest_from_paste(KOREAN_SAMPLE)
        self.assertIn("은혜", out)

    def test_empty_paste_raises(self) -> None:
        with self.assertRaises(IngestError):
            ingest_from_paste("   \n\n")


class IngestFromTxtTests(unittest.TestCase):
    def test_utf8_with_bom(self) -> None:
        data = b"\xef\xbb\xbf" + KOREAN_SAMPLE.encode("utf-8")
        out = ingest_from_txt(data)
        self.assertIn("은혜", out)

    def test_utf8_plain(self) -> None:
        out = ingest_from_txt(KOREAN_SAMPLE.encode("utf-8"))
        self.assertIn("은혜", out)

    def test_utf16(self) -> None:
        out = ingest_from_txt(KOREAN_SAMPLE.encode("utf-16"))
        self.assertIn("은혜", out)

    def test_empty_raises(self) -> None:
        with self.assertRaises(IngestError):
            ingest_from_txt(b"")

    def test_whitespace_only_raises(self) -> None:
        with self.assertRaises(IngestError):
            ingest_from_txt(b"   \n  \n")


class IngestFromDocxTests(unittest.TestCase):
    def test_reads_docx_paragraphs(self) -> None:
        from docx import Document

        doc = Document()
        doc.add_paragraph("오늘 우리는 하나님의 은혜를 보려고 합니다.")
        doc.add_paragraph("은혜는 단지 좋은 감정이 아닙니다.")
        buf = BytesIO()
        doc.save(buf)

        out = ingest_from_docx(buf.getvalue())
        self.assertIn("은혜", out)
        self.assertIn("감정", out)

    def test_empty_raises(self) -> None:
        with self.assertRaises(IngestError):
            ingest_from_docx(b"")

    def test_corrupt_docx_raises(self) -> None:
        with self.assertRaises(IngestError):
            ingest_from_docx(b"not a real docx file")

    def test_docx_with_no_text_raises(self) -> None:
        from docx import Document

        doc = Document()
        doc.add_paragraph("")
        doc.add_paragraph("   ")
        buf = BytesIO()
        doc.save(buf)
        with self.assertRaises(IngestError):
            ingest_from_docx(buf.getvalue())


class _MockGoogleDocsService:
    def __init__(self, *, document=None, raise_exc=None):
        self._document = document
        self._raise = raise_exc

    def documents(self):
        return self

    def get(self, *, documentId):
        self._last_doc_id = documentId
        return self

    def execute(self):
        if self._raise is not None:
            raise self._raise
        return self._document


class IngestFromGoogleDocsTests(unittest.TestCase):
    def test_extracts_text_from_doc(self) -> None:
        doc = {
            "body": {
                "content": [
                    {
                        "paragraph": {
                            "elements": [
                                {"textRun": {"content": "오늘 우리는 "}},
                                {"textRun": {"content": "은혜를 봅니다.\n"}},
                            ]
                        }
                    },
                    {
                        "paragraph": {
                            "elements": [
                                {"textRun": {"content": "기도합시다.\n"}},
                            ]
                        }
                    },
                ]
            }
        }
        service = _MockGoogleDocsService(document=doc)
        url = "https://docs.google.com/document/d/abc123XYZ/edit"
        out = ingest_from_google_docs(url, service)
        self.assertIn("은혜를", out)
        self.assertIn("기도합시다", out)
        self.assertEqual(service._last_doc_id, "abc123XYZ")

    def test_rejects_url_without_doc_id(self) -> None:
        service = _MockGoogleDocsService(document={"body": {"content": []}})
        with self.assertRaises(IngestError):
            ingest_from_google_docs(
                "https://example.com/foo", service
            )

    def test_propagates_api_failure_as_ingest_error(self) -> None:
        service = _MockGoogleDocsService(raise_exc=RuntimeError("403 quota"))
        with self.assertRaises(IngestError):
            ingest_from_google_docs(
                "https://docs.google.com/document/d/abc123XYZ/", service
            )

    def test_doc_with_no_text_raises(self) -> None:
        service = _MockGoogleDocsService(document={"body": {"content": []}})
        with self.assertRaises(IngestError):
            ingest_from_google_docs(
                "https://docs.google.com/document/d/abc123XYZ/", service
            )


class _StubTranslator:
    def __init__(self, *, prefix="[en]", raise_on=None):
        self._prefix = prefix
        self._raise_on = raise_on or set()
        self.calls: list[str] = []

    async def __call__(self, text: str) -> str:
        self.calls.append(text)
        if text in self._raise_on:
            raise RuntimeError("translator failed")
        return f"{self._prefix} {text}"


class BuildSermonTests(unittest.IsolatedAsyncioTestCase):
    async def test_happy_path_builds_sermon(self) -> None:
        translator = _StubTranslator()
        now = datetime(2026, 6, 16, tzinfo=timezone.utc)
        sermon = await build_sermon(
            sermonId="srm_abc",
            orgId="org_xyz",
            title="Test",
            sourceType="paste",
            sourceRef=None,
            text=KOREAN_SAMPLE,
            creatorUid="uid_owner",
            translator=translator,
            now=now,
        )
        self.assertEqual(sermon.sermonId, "srm_abc")
        self.assertEqual(sermon.orgId, "org_xyz")
        self.assertEqual(sermon.sourceType, "paste")
        self.assertEqual(len(sermon.segments), 3)
        self.assertEqual(sermon.createdAt, now)
        self.assertEqual(sermon.updatedAt, now)

    async def test_segment_ids_are_stable_and_zero_padded(self) -> None:
        translator = _StubTranslator()
        sermon = await build_sermon(
            sermonId="srm_abc",
            orgId="org_xyz",
            title="Test",
            sourceType="paste",
            sourceRef=None,
            text=KOREAN_SAMPLE,
            creatorUid="uid",
            translator=translator,
        )
        self.assertEqual(sermon.segments[0].segmentId, "S001")
        self.assertEqual(sermon.segments[1].segmentId, "S002")
        self.assertEqual(sermon.segments[2].segmentId, "S003")
        for i, seg in enumerate(sermon.segments, start=1):
            self.assertEqual(seg.order, i)

    async def test_reviewed_translation_prefilled_with_app_translation(self) -> None:
        translator = _StubTranslator(prefix="[t]")
        sermon = await build_sermon(
            sermonId="srm_abc",
            orgId="org_xyz",
            title="Test",
            sourceType="paste",
            sourceRef=None,
            text=KOREAN_SAMPLE,
            creatorUid="uid",
            translator=translator,
        )
        for seg in sermon.segments:
            self.assertEqual(seg.reviewedTranslation, seg.appTranslation)
            self.assertTrue(seg.appTranslation.startswith("[t]"))
            self.assertEqual(seg.status, "Draft")

    async def test_translation_failure_degrades_to_empty(self) -> None:
        # Per Design: a per-segment translation outage shouldn't lose the
        # whole sermon; the user can retranslate or edit.
        first_sentence = "오늘 우리는 하나님의 은혜를 보려고 합니다."
        translator = _StubTranslator(raise_on={first_sentence})
        sermon = await build_sermon(
            sermonId="srm_abc",
            orgId="org_xyz",
            title="Test",
            sourceType="paste",
            sourceRef=None,
            text=KOREAN_SAMPLE,
            creatorUid="uid",
            translator=translator,
        )
        self.assertEqual(sermon.segments[0].appTranslation, "")
        self.assertEqual(sermon.segments[0].reviewedTranslation, "")
        self.assertNotEqual(sermon.segments[1].appTranslation, "")

    async def test_empty_text_raises(self) -> None:
        translator = _StubTranslator()
        with self.assertRaises(IngestError):
            await build_sermon(
                sermonId="x",
                orgId="x",
                title="x",
                sourceType="paste",
                sourceRef=None,
                text="   ",
                creatorUid="x",
                translator=translator,
            )

    async def test_exceeds_max_segments_raises(self) -> None:
        translator = _StubTranslator()
        many = " ".join(["오늘 우리는 은혜를 봅니다."] * 50)
        with self.assertRaises(IngestError):
            await build_sermon(
                sermonId="x",
                orgId="x",
                title="x",
                sourceType="paste",
                sourceRef=None,
                text=many,
                creatorUid="x",
                translator=translator,
                max_segments=10,
            )


if __name__ == "__main__":
    unittest.main()
